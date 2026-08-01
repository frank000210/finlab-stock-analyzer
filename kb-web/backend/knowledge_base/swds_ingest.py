"""Ingest OR 作業處理流程 SWDS documents into a `swds_precedents` collection.

Goal (Phase 0 + Phase 1): turn ~258 historical SWDS (.doc/.docx, incl. OLE files
mislabeled as .docx) into structured precedent records so that a new SCN
description can retrieve similar past requirements together with the DB tables
(第2章 新增Table/欄位定義) and operation functions (第3章 相關程式配合) that were
actually changed. `suggest_swds` returns ranked precedents; `draft_swds` goes one
step further and renders a ready-to-edit SWDS skeleton (candidate 第2章 tables
aggregated by frequency + 第3章 functions grouped per closest precedent).

Extraction strategy (two readers, one shared parser):
  * Real OOXML (.docx, ZIP signature) -> python-docx, walking the body so
    paragraphs and tables keep document order.
  * Everything else (OLE .doc, or .docx that are really OLE binaries that Word
    refuses to re-encode) -> read tables/paragraphs directly through the Word
    COM object model. One Word instance serves every non-OOXML file.
  * Both readers emit generic blocks: ("para", text) and ("table", rows) where
    rows is List[List[str]]. `_parse_blocks` then classifies each table by its
    HEADER ROW (robust across analysts):
      - header has "Table名稱" + "欄位"           -> 第2章 table/欄位 definition
      - header has "程式名稱"                      -> 第3章 相關程式配合 (functions)
      - header has "Record內容"                    -> seed SQL / 申請事項 (records)
      - header has "修改內容"                      -> 版本異動說明 (changelog summary)
  * Parsed results are cached as JSON per file (keyed by path + mtime) so
    re-ingest does not re-open Word.
"""
from __future__ import annotations

import hashlib
import json
import os
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from pymongo.database import Database

from knowledge_base.documents import upsert_versioned

DEFAULT_ROOT = r"C:\Users\user\Documents\github\nextBSS\行動訂單業務分析文件\REF\OR作業處理流程"
COLLECTION = "swds_precedents"
KNOWLEDGE_SCOPE = "spec"
ZIP_SIGNATURE = b"PK\x03\x04"

Block = Tuple[str, Any]  # ("para", str) | ("table", List[List[str]])


# --------------------------------------------------------------------------- #
# discovery
# --------------------------------------------------------------------------- #
def discover_swds(root: Path) -> List[Path]:
    found: List[Path] = []
    for dirpath, _dirs, files in os.walk(root):
        for fn in files:
            low = fn.lower()
            if low.startswith("~$") or low.endswith(".partial"):
                continue
            if "swds" in low and (low.endswith(".doc") or low.endswith(".docx")):
                found.append(Path(dirpath) / fn)
    return sorted(found)


def _is_zip(path: Path) -> bool:
    try:
        with path.open("rb") as fh:
            return fh.read(4) == ZIP_SIGNATURE
    except OSError:
        return False


# --------------------------------------------------------------------------- #
# generic table/paragraph helpers
# --------------------------------------------------------------------------- #
def _collapse(cells: List[str]) -> List[str]:
    out: List[str] = []
    for c in cells:
        c = c.strip()
        if not out or out[-1] != c:  # collapse horizontally merged repeats
            out.append(c)
    return out


def _has_letter(text: str) -> bool:
    return any(c.isascii() and c.isalpha() for c in text)


def _header_kind(header: List[str]) -> Optional[str]:
    joined = " ".join(header)
    if "程式名稱" in joined:
        return "functions"
    if "Table名稱" in joined and "欄位" in joined:
        return "tables"
    if "Record內容" in joined:
        return "records"
    if "預期結果" in joined:
        return "test_cases"
    if "需求文件編號" in joined and "SWDS對應編號" in joined:
        return "trace_matrix"
    if "需求編號" in joined and "MR編號" in joined:
        return "trace_req"
    if "影響功能" in joined:
        return "impact"
    return None


def _clean_table_name(raw: str) -> Tuple[str, str, str]:
    """Return (action, ascii_identifier, chinese_desc) from a ch2 name cell."""
    action = ""
    text = raw.strip()
    for kw in ("新增表格", "修改表格", "異動表格", "新增", "修改", "異動"):
        if text.startswith(kw):
            action = "新增" if "新增" in kw else ("修改" if "修改" in kw else "異動")
            text = text[len(kw):].strip()
            break
    ident = ""
    rest = text
    buf = ""
    for ch in text:
        if ch.isascii() and (ch.isalnum() or ch == "_"):
            buf += ch
        else:
            if len(buf) >= 2:
                ident = buf
                rest = text[text.find(buf) + len(buf):].strip()
                break
            buf = ""
    if not ident and len(buf) >= 2:
        ident = buf
        rest = ""
    if ident.isdigit():  # pure dates/tags (e.g. 1060406) are not table names
        ident = ""
    chinese = rest.split("Dupls")[0].split("index")[0].strip(" ：:，,-")
    return action, ident, chinese


def _parse_tables_rows(rows: List[List[str]]) -> List[Dict[str, Any]]:
    grouped: Dict[str, Dict[str, Any]] = {}
    order: List[str] = []
    for vals in rows[1:]:
        if not vals or not vals[0].strip():
            continue
        name_cell = vals[0].strip()
        rest = vals[1:]
        db = ""
        cols = rest
        if rest and rest[0] and rest[0].lower().strip("[]").endswith("db"):
            db = rest[0].strip("[]")
            cols = rest[1:]
        is_subheader = bool(cols) and ("欄位名稱" in cols[0] or cols[0] == "欄位")
        if name_cell not in grouped:
            action, ident, chinese = _clean_table_name(name_cell)
            grouped[name_cell] = {
                "action": action,
                "name": ident or chinese or name_cell[:40],
                "db": db,
                "chinese": chinese,
                "columns": [],
            }
            order.append(name_cell)
        elif db and not grouped[name_cell]["db"]:
            grouped[name_cell]["db"] = db
        if is_subheader:
            continue
        col = {
            "name": cols[0] if len(cols) > 0 else "",
            "type": cols[1] if len(cols) > 1 else "",
            "desc": cols[2] if len(cols) > 2 else "",
        }
        if col["name"]:
            grouped[name_cell]["columns"].append(col)
    # keep only real DB tables (ascii-letter identifier); drop date/tag noise rows
    return [grouped[k] for k in order if _has_letter(grouped[k]["name"])]


def _parse_functions_rows(rows: List[List[str]]) -> List[Dict[str, str]]:
    out: List[Dict[str, str]] = []
    group = ""
    for vals in rows[1:]:
        if not vals or not vals[0].strip():
            continue
        name = vals[0].strip()
        note = vals[1].strip() if len(vals) > 1 else ""
        if not note and len(name) <= 6:  # 售前/售中/售後 group label row
            group = name
            continue
        out.append({"group": group, "name": name[:200], "note": note[:400]})
    return out


def _multi(value: str, sep: str = "；") -> str:
    """Join a multi-line cell into a single line, preserving item separation."""
    parts = [p.strip() for p in str(value or "").replace("\r", "\n").split("\n")]
    return sep.join(p for p in parts if p)


def _parse_test_cases_rows(rows: List[List[str]]) -> List[Dict[str, str]]:
    """Parse a 測試案例 table → {name,precondition,operation,method,expected,note}."""
    header = rows[0]
    idx: Dict[str, int] = {}
    for i, h in enumerate(header):
        hn = _norm(h)
        if "案例名稱" in hn:
            idx.setdefault("name", i)
        elif "前提條件" in hn:
            idx.setdefault("precondition", i)
        elif "受理作業" in hn:
            idx.setdefault("operation", i)
        elif "操作方式" in hn:
            idx.setdefault("method", i)
        elif "預期結果" in hn:
            idx.setdefault("expected", i)
        elif "備註" in hn:
            idx.setdefault("note", i)
    out: List[Dict[str, str]] = []
    for vals in rows[1:]:
        if not any((v or "").strip() for v in vals):
            continue

        def g(key: str) -> str:
            i = idx.get(key)
            return _norm(vals[i]) if i is not None and i < len(vals) else ""

        def gm(key: str) -> str:
            i = idx.get(key)
            return _multi(vals[i]) if i is not None and i < len(vals) else ""

        rec = {
            "name": g("name")[:120],
            "precondition": gm("precondition")[:300],
            "operation": g("operation")[:200],
            "method": gm("method")[:300],
            "expected": gm("expected")[:400],
            "note": g("note")[:200],
        }
        if not (rec["name"] or rec["expected"] or rec["precondition"]):
            continue
        if rec["name"] == "案例名稱" or rec["expected"] == "預期結果":
            continue
        out.append(rec)
    return out


def _parse_trace_req_rows(rows: List[List[str]]) -> Dict[str, str]:
    """Parse a 需求提案 key-value table → {req_no, mr_no, req_name}."""
    out = {"req_no": "", "mr_no": "", "req_name": ""}
    for vals in rows:
        for i, cell in enumerate(vals):
            c = _norm(cell)
            nxt = _norm(vals[i + 1]) if i + 1 < len(vals) else ""
            if not nxt:
                continue
            if c == "需求編號" and not out["req_no"]:
                out["req_no"] = nxt[:60]
            elif "MR編號" in c and not out["mr_no"]:
                out["mr_no"] = nxt[:60]
            elif c in ("需求名稱", "需求主旨", "需求標題", "需求摘要") and not out["req_name"]:
                out["req_name"] = nxt[:200]
    return out


def _parse_trace_matrix_rows(rows: List[List[str]]) -> List[Dict[str, str]]:
    """Parse a 需求文件編號↔SWDS對應編號 traceability matrix."""
    out: List[Dict[str, str]] = []
    for vals in rows[1:]:
        if not vals or not (vals[0] or "").strip():
            continue
        req_doc = _norm(vals[0])
        swds_code = _norm(vals[1]) if len(vals) > 1 else ""
        if req_doc == "需求文件編號":
            continue
        if req_doc or swds_code:
            out.append({"req_doc": req_doc[:80], "swds_code": swds_code[:80]})
    return out


def _parse_impact_rows(rows: List[List[str]]) -> List[Dict[str, str]]:
    """Parse an 影響功能 table (Tag/主要異動/異動說明/影響功能)."""
    header = rows[0]
    idx: Dict[str, int] = {}
    for i, h in enumerate(header):
        hn = _norm(h)
        if hn in ("Tag", "TAG"):
            idx.setdefault("tag", i)
        if "主要異動" in hn:
            idx.setdefault("change", i)
        if "異動說明" in hn:
            idx.setdefault("desc", i)
        if "影響功能" in hn:
            idx.setdefault("functions", i)
    out: List[Dict[str, str]] = []
    for vals in rows[1:]:
        if not any((v or "").strip() for v in vals):
            continue

        def g(key: str) -> str:
            i = idx.get(key)
            return _norm(vals[i]) if i is not None and i < len(vals) else ""

        def gm(key: str) -> str:
            i = idx.get(key)
            return _multi(vals[i]) if i is not None and i < len(vals) else ""

        rec = {
            "tag": g("tag")[:60],
            "change": g("change")[:120],
            "desc": gm("desc")[:300],
            "functions": gm("functions")[:400],
        }
        if rec["functions"] == "影響功能":
            continue
        if rec["functions"] or rec["change"] or rec["desc"]:
            out.append(rec)
    return out


def _impact_func_list(impact: List[Dict[str, str]]) -> List[str]:
    out: List[str] = []
    for it in impact:
        for f in str(it.get("functions", "")).split("；"):
            f = f.strip()
            if f and f not in out:
                out.append(f)
    return out[:50]


def _parse_blocks(blocks: List[Block]) -> Dict[str, Any]:
    req_lines: List[str] = []
    changelog_lines: List[str] = []
    ch2: List[Dict[str, Any]] = []
    ch3: List[Dict[str, str]] = []
    records: List[str] = []
    test_cases: List[Dict[str, str]] = []
    trace: Dict[str, Any] = {"req_no": "", "mr_no": "", "req_name": "", "doc_map": []}
    impact: List[Dict[str, str]] = []
    seen_analysis = False

    for kind, payload in blocks:
        if kind == "para":
            t = str(payload).strip()
            if not t:
                continue
            if not seen_analysis and len(t) >= 6 and not t.endswith("："):
                if not any(t.startswith(h) for h in (
                    "版本", "需求提案", "影響範圍", "SWDS", "新增Table", "相關程式",
                )):
                    req_lines.append(t)
            continue
        rows = payload
        if not rows:
            continue
        header = rows[0]
        kindt = _header_kind(header)
        if kindt == "tables":
            seen_analysis = True
            ch2.extend(_parse_tables_rows(rows))
        elif kindt == "functions":
            seen_analysis = True
            ch3.extend(_parse_functions_rows(rows))
        elif kindt == "records":
            seen_analysis = True
            for vals in rows[1:]:
                if vals and vals[0].strip():
                    records.append(vals[0].strip())
        elif kindt == "test_cases":
            seen_analysis = True
            test_cases.extend(_parse_test_cases_rows(rows))
        elif kindt == "trace_matrix":
            trace["doc_map"].extend(_parse_trace_matrix_rows(rows))
        elif kindt == "trace_req":
            tr = _parse_trace_req_rows(rows)
            for k in ("req_no", "mr_no", "req_name"):
                if tr.get(k) and not trace[k]:
                    trace[k] = tr[k]
        elif kindt == "impact":
            seen_analysis = True
            impact.extend(_parse_impact_rows(rows))
        elif "修改內容" in " ".join(header):
            try:
                ci = header.index("修改內容")
            except ValueError:
                ci = None
            if ci is not None:
                for vals in rows[1:]:
                    if len(vals) > ci:
                        v = vals[ci].strip()
                        if v and v not in ("初版", "修改內容") and len(v) >= 4:
                            changelog_lines.append(v)

    parts: List[str] = []
    if req_lines:
        parts.append("\n".join(req_lines))
    if changelog_lines:
        parts.append("【版本異動摘要】" + "；".join(changelog_lines))
    requirement_text = "\n".join(parts)[:4000].strip()
    return {
        "requirement_text": requirement_text,
        "tables": ch2,
        "functions": ch3,
        "seed_records": records,
        "test_cases": test_cases,
        "trace": trace,
        "impact": impact,
    }


# --------------------------------------------------------------------------- #
# readers
# --------------------------------------------------------------------------- #
def _blocks_from_docx(path: Path) -> List[Block]:
    from docx import Document
    from docx.oxml.ns import qn

    d = Document(str(path))
    tables = {id(t._tbl): t for t in d.tables}
    paras = {id(p._p): p for p in d.paragraphs}
    blocks: List[Block] = []
    for child in d.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = paras.get(id(child))
            if p is not None and p.text.strip():
                blocks.append(("para", p.text.strip()))
        elif child.tag == qn("w:tbl"):
            t = tables.get(id(child))
            if t is None:
                continue
            rows: List[List[str]] = []
            for row in t.rows:
                cells = [
                    " ".join(pp.text.strip() for pp in c.paragraphs if pp.text.strip()).strip()
                    for c in row.cells
                ]
                rows.append(_collapse(cells))
            blocks.append(("table", rows))
    return blocks


def _clean_com_text(text: str) -> str:
    return text.replace("\r", " ").replace("\x07", " ").replace("\x0b", " ").strip()


def _read_com_table(table) -> List[List[str]]:
    rows: List[List[str]] = []
    try:
        row_iter = list(table.Rows)
    except Exception:  # noqa: BLE001 - vertically merged tables cannot enumerate rows
        return rows
    for row in row_iter:
        try:
            cells = [_clean_com_text(c.Range.Text) for c in row.Cells]
        except Exception:  # noqa: BLE001
            continue
        rows.append(_collapse(cells))
    return rows


def _blocks_from_com(word, path: Path) -> List[Block]:
    doc = word.Documents.Open(str(path), False, True)
    try:
        table_objs = list(doc.Tables)
        first_analysis_start: Optional[int] = None
        table_blocks: List[Tuple[int, Block]] = []
        for t in table_objs:
            rows = _read_com_table(t)
            start = int(t.Range.Start)
            table_blocks.append((start, ("table", rows)))
            if rows and _header_kind(rows[0]) is not None and first_analysis_start is None:
                first_analysis_start = start
        # requirement paragraphs = text before the first analysis table
        pre_lines: List[str] = []
        if first_analysis_start and first_analysis_start > 0:
            pre = _clean_com_text(doc.Range(0, first_analysis_start).Text)
            for line in pre.split("\n"):
                line = line.strip()
                if line:
                    pre_lines.append(line)
        blocks: List[Block] = [("para", ln) for ln in pre_lines]
        for _start, blk in sorted(table_blocks, key=lambda x: x[0]):
            blocks.append(blk)
        return blocks
    finally:
        doc.Close(False)


# --------------------------------------------------------------------------- #
# extraction with per-file JSON cache
# --------------------------------------------------------------------------- #
def _cache_json_path(cache_dir: Path, src: Path) -> Path:
    digest = hashlib.sha1(str(src.resolve()).encode("utf-8")).hexdigest()[:16]
    return cache_dir / f"{digest}.json"


def extract_swds(src: Path, word, cache_dir: Path) -> Tuple[Dict[str, Any], str]:
    """Return (parsed, source_format). Uses/refreshes a JSON cache keyed by mtime."""
    cache = _cache_json_path(cache_dir, src)
    mtime = src.stat().st_mtime
    if cache.exists():
        try:
            payload = json.loads(cache.read_text(encoding="utf-8"))
            if payload.get("mtime") == mtime:
                return payload["parsed"], payload.get("format", "cache")
        except (OSError, ValueError, KeyError):
            pass

    if _is_zip(src):
        blocks = _blocks_from_docx(src)
        fmt = "docx"
    else:
        blocks = _blocks_from_com(word, src)
        fmt = "ole-com"
    parsed = _parse_blocks(blocks)
    cache.write_text(
        json.dumps({"mtime": mtime, "format": fmt, "parsed": parsed}, ensure_ascii=False),
        encoding="utf-8",
    )
    return parsed, fmt


# --------------------------------------------------------------------------- #
# record assembly
# --------------------------------------------------------------------------- #
def _slug(root: Path, path: Path) -> str:
    rel = str(path.resolve()).replace(str(root.resolve()), "").strip("\\/")
    digest = hashlib.sha1(rel.encode("utf-8")).hexdigest()[:12]
    return f"swds-{digest}"


def _category(root: Path, path: Path) -> str:
    try:
        rel = path.resolve().relative_to(root.resolve())
        return rel.parts[0] if len(rel.parts) > 1 else ""
    except ValueError:
        return ""


def _render_content(title: str, category: str, parsed: Dict[str, Any]) -> str:
    lines = [f"# {title}", ""]
    if category:
        lines += [f"**領域分類**：{category}", ""]
    req = parsed.get("requirement_text", "")
    if req:
        lines += ["## 需求說明（SCN）", "", req, ""]
    tables = parsed.get("tables", [])
    lines += [f"## 第2章 新增/異動 Table（{len(tables)}）", ""]
    if tables:
        lines += ["| 動作 | Table | DB | 說明 | 欄位數 |", "|---|---|---|---|---|"]
        for t in tables:
            lines.append(
                "| {action} | {name} | {db} | {ch} | {n} |".format(
                    action=t.get("action", ""), name=t.get("name", ""),
                    db=t.get("db", ""), ch=t.get("chinese", ""),
                    n=len(t.get("columns", [])),
                )
            )
        lines.append("")
    funcs = parsed.get("functions", [])
    lines += [f"## 第3章 相關程式配合（{len(funcs)}）", ""]
    if funcs:
        lines += ["| 售別 | 程式名稱 | 內容摘要 |", "|---|---|---|"]
        for f in funcs:
            note = (f.get("note", "") or "").replace("|", "／").replace("\n", " ")[:80]
            lines.append(
                "| {g} | {name} | {note} |".format(
                    g=f.get("group", ""),
                    name=(f.get("name", "") or "").replace("|", "／"),
                    note=note,
                )
            )
        lines.append("")
    impact = parsed.get("impact", [])
    if impact:
        lines += [f"## 影響功能（{len(impact)}）", ""]
        lines += ["| Tag | 主要異動 | 影響功能 |", "|---|---|---|"]
        for it in impact:
            lines.append(
                "| {t} | {c} | {f} |".format(
                    t=_md_cell(it.get("tag", "")),
                    c=_md_cell(it.get("change", "")),
                    f=_md_cell(it.get("functions", "")),
                )
            )
        lines.append("")
    tcs = parsed.get("test_cases", [])
    if tcs:
        lines += [f"## 測試案例（{len(tcs)}）", ""]
        lines += [
            "| 案例名稱 | 前提條件 | 受理作業 | 操作方式 | 預期結果 |",
            "|---|---|---|---|---|",
        ]
        for c in tcs:
            lines.append(
                "| {n} | {p} | {o} | {m} | {e} |".format(
                    n=_md_cell(c.get("name", "")),
                    p=_md_cell(c.get("precondition", "")),
                    o=_md_cell(c.get("operation", "")),
                    m=_md_cell(c.get("method", "")),
                    e=_md_cell(c.get("expected", "")),
                )
            )
        lines.append("")
    trace = parsed.get("trace", {}) or {}
    dm = trace.get("doc_map", [])
    if trace.get("req_no") or trace.get("mr_no") or trace.get("req_name") or dm:
        lines += ["## 追溯資訊", ""]
        if trace.get("req_no"):
            lines.append(f"- 需求編號：{trace['req_no']}")
        if trace.get("mr_no"):
            lines.append(f"- MR編號：{trace['mr_no']}")
        if trace.get("req_name"):
            lines.append(f"- 需求名稱：{trace['req_name']}")
        if dm:
            lines += ["", "| 需求文件編號 | SWDS對應編號 |", "|---|---|"]
            for d in dm:
                lines.append(
                    f"| {_md_cell(d.get('req_doc', ''))} | {_md_cell(d.get('swds_code', ''))} |"
                )
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def _summary(parsed: Dict[str, Any]) -> str:
    tnames = "、".join(t.get("name", "") for t in parsed.get("tables", [])[:6] if t.get("name"))
    nfun = len(parsed.get("functions", []))
    req = (parsed.get("requirement_text", "") or "").splitlines()
    head = req[0][:60] if req else ""
    parts = []
    if head:
        parts.append(head)
    if tnames:
        parts.append(f"改表:{tnames}")
    parts.append(f"配合功能{nfun}項")
    ntc = len(parsed.get("test_cases", []))
    if ntc:
        parts.append(f"測試案例{ntc}則")
    return "；".join(parts)[:200]


def build_record(root: Path, path: Path, parsed: Dict[str, Any], fmt: str) -> Dict[str, Any]:
    title = path.stem
    category = _category(root, path)
    cat_tag = category.split("_", 1)[-1] if "_" in category else category
    tags = ["SWDS"]
    if cat_tag:
        tags.append(cat_tag)
    return {
        "slug": _slug(root, path),
        "title": title,
        "category": category,
        "source_path": str(path),
        "source_format": fmt,
        "requirement_text": parsed.get("requirement_text", ""),
        "tables": parsed.get("tables", []),
        "functions": parsed.get("functions", []),
        "seed_records": parsed.get("seed_records", []),
        "table_names": [t.get("name", "") for t in parsed.get("tables", []) if t.get("name")],
        "function_names": [f.get("name", "") for f in parsed.get("functions", [])],
        "test_cases": parsed.get("test_cases", []),
        "trace": parsed.get("trace", {}),
        "impact": parsed.get("impact", []),
        "test_case_names": [c.get("name", "") for c in parsed.get("test_cases", []) if c.get("name")],
        "impact_functions": _impact_func_list(parsed.get("impact", [])),
        "requirement_no": (parsed.get("trace", {}) or {}).get("req_no", ""),
        "mr_no": (parsed.get("trace", {}) or {}).get("mr_no", ""),
        "summary": _summary(parsed),
        "tags": tags,
        "content": _render_content(title, category, parsed),
        "knowledge_scope": KNOWLEDGE_SCOPE,
    }


# --------------------------------------------------------------------------- #
# orchestration
# --------------------------------------------------------------------------- #
def ingest_swds(
    db: Database,
    root: Optional[Path] = None,
    apply: bool = False,
    limit: Optional[int] = None,
    cache_dir: Optional[Path] = None,
) -> Dict[str, Any]:
    root = Path(root) if root else Path(DEFAULT_ROOT)
    cache_dir = Path(cache_dir) if cache_dir else Path(tempfile.gettempdir()) / "kb_swds_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)

    files = discover_swds(root)
    if limit:
        files = files[:limit]

    needs_word = any(not _is_zip(p) for p in files)
    word = None
    if needs_word:
        try:
            import win32com.client  # type: ignore
        except ModuleNotFoundError as exc:  # pragma: no cover - windows only
            raise RuntimeError("pywin32 is required to read legacy Word files.") from exc
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

    parse_failures: List[Dict[str, str]] = []
    upserts = 0
    empty = 0
    fmt_counts: Dict[str, int] = {}
    samples: List[Dict[str, Any]] = []

    try:
        for i, src in enumerate(files, 1):
            try:
                parsed, fmt = extract_swds(src, word, cache_dir)
            except Exception as exc:  # noqa: BLE001
                parse_failures.append({"file": str(src), "error": f"{exc!r}"[:200]})
                continue
            fmt_counts[fmt] = fmt_counts.get(fmt, 0) + 1
            if not parsed["tables"] and not parsed["functions"]:
                empty += 1
            record = build_record(root, src, parsed, fmt)
            if len(samples) < 6:
                samples.append({
                    "title": record["title"],
                    "category": record["category"],
                    "format": fmt,
                    "tables": record["table_names"][:8],
                    "functions_count": len(record["function_names"]),
                })
            if apply:
                payload = {k: v for k, v in record.items() if k != "slug"}
                upsert_versioned(db, COLLECTION, "slug", record["slug"], payload)
                upserts += 1
            if i % 40 == 0:
                print(f"  processed {i}/{len(files)}", flush=True)
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:  # pragma: no cover
                pass

    return {
        "apply": apply,
        "root": str(root),
        "cache_dir": str(cache_dir),
        "discovered": len(files),
        "upserted": upserts,
        "empty_sections": empty,
        "format_counts": fmt_counts,
        "parse_failures": parse_failures,
        "samples": samples,
    }


def suggest_swds(db: Database, text: str, top_k: int = 5) -> Dict[str, Any]:
    """Given an SCN description, return similar SWDS precedents with the DB
    tables and operation functions they changed."""
    from knowledge_base.semantic_search import semantic_search

    hits = semantic_search(
        db, text, top_k=top_k, include_collections=[COLLECTION], scope="all"
    )
    results: List[Dict[str, Any]] = []
    table_counter: Dict[str, int] = {}
    function_counter: Dict[str, int] = {}
    for hit in hits:
        doc = hit.get("doc", {})
        tables = doc.get("tables", [])
        functions = doc.get("functions", [])
        for t in tables:
            name = t.get("name", "")
            if name:
                table_counter[name] = table_counter.get(name, 0) + 1
        for f in functions:
            name = f.get("name", "")
            if name:
                function_counter[name] = function_counter.get(name, 0) + 1
        results.append({
            "slug": doc.get("slug", ""),
            "title": doc.get("title", ""),
            "category": doc.get("category", ""),
            "score": round(float(hit.get("semantic_score", 0.0)), 4),
            "summary": doc.get("summary", ""),
            "tables": [
                {"action": t.get("action", ""), "name": t.get("name", ""),
                 "db": t.get("db", ""), "chinese": t.get("chinese", "")}
                for t in tables
            ],
            "functions": [
                {"group": f.get("group", ""), "name": f.get("name", "")}
                for f in functions
            ],
        })
    return {
        "query": text,
        "matches": len(results),
        "results": results,
        "aggregated_tables": sorted(
            table_counter.items(), key=lambda kv: kv[1], reverse=True
        ),
        "aggregated_functions": sorted(
            function_counter.items(), key=lambda kv: kv[1], reverse=True
        )[:20],
    }


def _norm(value: str) -> str:
    """Collapse whitespace/newlines in a string to a single line."""
    return " ".join(str(value or "").split())


def _md_cell(value: str) -> str:
    """Make a string safe for a markdown table cell."""
    return _norm(value).replace("|", "/") or " "


def _strip_swds_prefix(title: str) -> str:
    t = str(title or "").strip()
    for p in ("SWDS_", "SWDS-", "SWDS"):
        if t.startswith(p):
            return t[len(p):].strip() or t
    return t


def _fmt_sources(titles: List[str], limit: int = 3) -> str:
    seen: List[str] = []
    for t in titles:
        d = _strip_swds_prefix(t)
        if d and d not in seen:
            seen.append(d)
    if not seen:
        return " "
    shown = "、".join(seen[:limit])
    if len(seen) > limit:
        shown += f" 等{len(seen)}份"
    return shown


# Function group ordering for the drafted 第3章
_GROUP_ORDER = ["售前", "售中", "售後"]


def _resolve_pages(
    db: Database, func_names: List[str], limit_per: int = 3
) -> Dict[str, List[Dict[str, Any]]]:
    """Match SWDS function names to page entries.

    Function names in SWDS are often menu navigation paths with multiple sub-items
    (e.g., "3.1 輸入門號 -個案增刪優惠 -大批增優惠"). We split by common delimiters
    and try each CJK segment as a search token against code_e2e_chain (ORW) first,
    then code_menu (MBMS) if no ORW hits.

    Returns {func_name: [{"source", "menu_name", "base_code", "channel",
                          "backing_bean", "delegates", "bd_calls", "action_methods",
                          "menu_path", "jsp_path"}]}
    """
    import re as _re

    def _cjk_tokens(text: str) -> List[str]:
        """Extract distinct 3-6 char CJK substrings from a menu-path-style string."""
        # split on: ( ) ( ) [ ] - ＋ + space newline , 、；。 and digit-sequences
        parts = _re.split(r"[\s\(\)\（\）\[\]＋\+\-\,、；。\n]|[0-9]+\.", text)
        tokens: List[str] = []
        seen: set = set()
        for part in parts:
            cjk = "".join(c for c in part if "\u4e00" <= c <= "\u9fff")
            # try 4-char and 3-char windows
            for start in range(0, max(1, len(cjk) - 2)):
                for wlen in (4, 3):
                    tok = cjk[start: start + wlen]
                    if len(tok) >= 3 and tok not in seen:
                        tokens.append(tok)
                        seen.add(tok)
        return tokens[:20]  # cap to avoid too many DB queries

    result: Dict[str, List[Dict[str, Any]]] = {}
    for fname in func_names:
        tokens = _cjk_tokens(fname)
        if not tokens:
            continue
        matches: List[Dict[str, Any]] = []
        seen_bc: set = set()

        # Source 1: ORW code_e2e_chain — try each token until we have enough matches
        for tok in tokens:
            if len(matches) >= limit_per:
                break
            for doc in db.code_e2e_chain.find(
                {"content": {"$regex": tok, "$options": "i"}},
                {
                    "base_code": 1, "content": 1, "channel": 1, "menu_path_str": 1,
                    "jsp_path": 1, "backing_bean": 1, "backing_bean_short": 1,
                    "delegates": 1, "bd_calls": 1, "action_methods": 1,
                },
                limit=limit_per,
            ):
                bc = doc.get("base_code", "")
                if bc in seen_bc:
                    continue
                seen_bc.add(bc)
                matches.append(
                    {
                        "source": "orw",
                        "menu_name": doc.get("content", ""),
                        "base_code": bc,
                        "channel": doc.get("channel", "ORW"),
                        "backing_bean": doc.get("backing_bean") or doc.get("backing_bean_short", ""),
                        "delegates": doc.get("delegates") or [],
                        "bd_calls": doc.get("bd_calls") or {},
                        "action_methods": (doc.get("action_methods") or [])[:4],
                        "menu_path": doc.get("menu_path_str", ""),
                        "jsp_path": doc.get("jsp_path", ""),
                    }
                )

        # Source 2: MBMS code_menu — fill remaining slots with leaf pages
        if len(matches) < limit_per:
            for tok in tokens[:4]:
                if len(matches) >= limit_per:
                    break
                for doc in db.code_menu.find(
                    {
                        "$or": [
                            {"menu_path_str": {"$regex": tok, "$options": "i"}},
                            {"content": {"$regex": tok, "$options": "i"}},
                        ],
                        "is_leaf": True,
                        "channel": "MBMS",
                    },
                    {"content": 1, "base_code": 1, "channel": 1, "menu_path_str": 1, "jsp_path": 1},
                    limit=limit_per,
                ):
                    bc = doc.get("base_code", "")
                    if bc in seen_bc:
                        continue
                    seen_bc.add(bc)
                    matches.append(
                        {
                            "source": "mbms",
                            "menu_name": doc.get("content", ""),
                            "base_code": bc,
                            "channel": doc.get("channel", "MBMS"),
                            "backing_bean": "",
                            "delegates": [],
                            "bd_calls": {},
                            "action_methods": [],
                            "menu_path": doc.get("menu_path_str", ""),
                            "jsp_path": doc.get("jsp_path", ""),
                        }
                    )

        if matches:
            result[fname] = matches
    return result


def _resolve_program_chain(
    db: Database, base_codes: List[str]
) -> Dict[str, Dict[str, Any]]:
    """For each base_code, return BackingBean + BD delegate + action method chain.
    Used when base_codes come from code_e2e_chain (ORW system).
    """
    result: Dict[str, Dict[str, Any]] = {}
    for code in base_codes:
        chain = db.code_e2e_chain.find_one({"base_code": code})
        if not chain:
            continue
        delegates = chain.get("delegates") or []
        bd_calls = chain.get("bd_calls") or {}
        bd_methods: List[str] = []
        for bd in delegates[:3]:
            for m in (bd_calls.get(bd) or [])[:4]:
                bd_methods.append(f"{bd}.{m}")
        bb = chain.get("backing_bean") or chain.get("backing_bean_short", "")
        java = chain.get("java_source_path", "")
        if java and "MOrderSource/" in java:
            java = java.split("MOrderSource/")[-1]
        result[code] = {
            "backing_bean": bb,
            "delegates": delegates,
            "bd_methods": bd_methods[:6],
            "action_methods": (chain.get("action_methods") or [])[:4],
            "java_source_path": java,
        }
    return result


def draft_swds(
    db: Database,
    text: str,
    top_k: int = 8,
    title: Optional[str] = None,
    category: Optional[str] = None,
) -> Dict[str, Any]:
    """Given an SCN description, produce a ready-to-edit SWDS skeleton.

    Retrieves the most similar precedents and aggregates the DB tables (第2章)
    and operation functions (第3章) they changed, annotating each candidate with
    how many precedents used it (frequency = confidence) and which precedents
    they came from. The output is a markdown draft the analyst edits down to the
    tables/functions that truly apply to the new SCN.
    """
    from knowledge_base.semantic_search import semantic_search

    hits = semantic_search(
        db, text, top_k=top_k, include_collections=[COLLECTION], scope="all"
    )

    table_agg: Dict[str, Dict[str, Any]] = {}
    func_agg: Dict[str, Dict[str, Any]] = {}
    precedents: List[Dict[str, Any]] = []

    for hit in hits:
        doc = hit.get("doc", {})
        dtitle = doc.get("title", "")
        score = round(float(hit.get("semantic_score", 0.0)), 4)
        precedents.append({
            "slug": doc.get("slug", ""),
            "title": dtitle,
            "category": doc.get("category", ""),
            "score": score,
            "summary": doc.get("summary", ""),
            "functions": doc.get("functions", []) or [],
            "test_cases": doc.get("test_cases", []) or [],
        })
        for t in doc.get("tables", []):
            name = t.get("name", "")
            if not name:
                continue
            e = table_agg.setdefault(name, {
                "actions": [], "dbs": [], "chineses": [], "cols": 0, "sources": [],
            })
            if t.get("action"):
                e["actions"].append(t["action"])
            if t.get("db"):
                e["dbs"].append(t["db"])
            if t.get("chinese"):
                e["chineses"].append(t["chinese"])
            e["cols"] = max(e["cols"], len(t.get("columns", []) or []))
            e["sources"].append(dtitle)
        for f in doc.get("functions", []):
            name = f.get("name", "")
            if not name:
                continue
            e = func_agg.setdefault(name, {"group": f.get("group", ""), "sources": [], "notes": []})
            if not e["group"] and f.get("group"):
                e["group"] = f.get("group")
            e["sources"].append(dtitle)
            if f.get("note"):
                e["notes"].append(f.get("note"))

    def _distinct(seq: List[str]) -> List[str]:
        out: List[str] = []
        for x in seq:
            if x and x not in out:
                out.append(x)
        return out

    # ---- build ordered candidate lists (by precedent frequency desc) ----
    table_rows = []
    for name, e in table_agg.items():
        freq = len(set(e["sources"]))
        table_rows.append((freq, name, e))
    table_rows.sort(key=lambda r: (-r[0], r[1]))

    func_rows = []
    for name, e in func_agg.items():
        freq = len(set(e["sources"]))
        func_rows.append((freq, name, e))
    func_rows.sort(key=lambda r: (-r[0], r[1]))

    # ---- Phase 1+2: resolve page names + program chains ----
    # Collect distinct function names from the top-4 closest precedents.
    detail_n_pre = min(4, len(precedents))
    all_func_names: List[str] = []
    seen_fn: set = set()
    for p in precedents[:detail_n_pre]:
        for f in p.get("functions", []):
            n = _norm(f.get("name", ""))
            if n and n not in seen_fn:
                all_func_names.append(n)
                seen_fn.add(n)

    page_map = _resolve_pages(db, all_func_names)          # func_name -> [page entries]
    all_base_codes: List[str] = []
    seen_bc: set = set()
    for pages in page_map.values():
        for pg in pages:
            bc = pg.get("base_code", "")
            if bc and bc not in seen_bc:
                all_base_codes.append(bc)
                seen_bc.add(bc)
    chain_map = _resolve_program_chain(db, all_base_codes)  # base_code -> chain

    # ---- render markdown skeleton ----
    disp_title = title or "（待填：新需求標題）"
    disp_cat = category or (precedents[0]["category"] if precedents else "（待填）")
    lines: List[str] = []
    lines.append(f"# SWDS_{disp_title}（草稿）")
    lines.append("")
    lines.append(f"> 本草稿由 `swds-draft` 依 SCN 檢索 {len(precedents)} 份相似前例自動彙整，")
    lines.append("> **候選項需人工比對本次 SCN 後增刪**；前例數越高代表越多相似需求動過，優先檢視。")
    lines.append("")
    lines.append(f"**領域分類**：{disp_cat}")
    lines.append("")
    lines.append("## 需求說明（SCN）")
    lines.append("")
    lines.append(_norm(text) or "（待填）")
    lines.append("")

    # 第2章
    lines.append(f"## 第2章 候選新增/異動 Table（{len(table_rows)}）")
    lines.append("")
    if table_rows:
        lines.append("| 建議動作 | Table | DB | 說明 | 參考欄位數 | 前例數 | 來源前例 |")
        lines.append("|---|---|---|---|---|---|---|")
        for freq, name, e in table_rows:
            action = "/".join(_distinct(e["actions"])) or " "
            db_name = "/".join(_distinct(e["dbs"])) or " "
            chinese = "；".join(_distinct(e["chineses"])[:2]) or " "
            cols = e["cols"] or " "
            lines.append(
                f"| {_md_cell(action)} | {_md_cell(name)} | {_md_cell(db_name)} | "
                f"{_md_cell(chinese)} | {cols} | {freq} | {_md_cell(_fmt_sources(e['sources']))} |"
            )
    else:
        lines.append("_（前例未偵測到第2章 Table 異動；請確認是否為純功能調整需求）_")
    lines.append("")

    # 第3章 — grouped BY PRECEDENT. Function names carry menu paths + date tags
    # so they rarely recur across precedents; a flat frequency list is noise.
    # Per-precedent detail (top N closest) mirrors how an analyst browses, plus a
    # small "recurring across precedents" callout for the rare high-signal repeats.
    recurring = [(freq, name, e) for freq, name, e in func_rows if freq >= 2]
    detail_n = min(4, len(precedents))
    lines.append(
        f"## 第3章 候選相關程式配合（共 {len(func_rows)} 項，以下列前 {detail_n} 份前例明細）"
    )
    lines.append("")
    if recurring:
        lines.append(f"**跨前例重複出現的功能（{len(recurring)} 項，高機率必改）**")
        lines.append("")
        for freq, name, e in recurring:
            lines.append(f"- **{name}**（{freq} 份前例：{_fmt_sources(e['sources'])}）")
        lines.append("")
    if func_rows:
        lines.append("### 各前例功能配合明細")
        lines.append("")
        for p in precedents[:detail_n]:
            fns = p.get("functions", [])
            if not fns:
                continue
            lines.append(f"#### {_strip_swds_prefix(p['title'])}（相似度 {p['score']}）")
            lines.append("")
            pg: Dict[str, List[Dict[str, Any]]] = {}
            for f in fns:
                pg.setdefault(f.get("group") or "其他", []).append(f)
            pgroups = [g for g in _GROUP_ORDER if g in pg] + [
                g for g in pg if g not in _GROUP_ORDER
            ]
            for g in pgroups:
                lines.append(f"- **{g}**")
                for f in pg[g]:
                    fname = _norm(f.get("name", ""))
                    if not fname:
                        continue
                    note = _norm(f.get("note", ""))
                    if len(note) > 60:
                        note = note[:60] + "…"
                    lines.append(f"  - {fname}" + (f" — {note}" if note else ""))
            lines.append("")
    else:
        lines.append("_（前例未偵測到第3章功能配合）_")
        lines.append("")

    # ---- 第3章 頁面 ↔ 程式鏈對應（Phase 1 + Phase 2）----
    lines.append("## 第3章 頁面名稱 ↔ 程式鏈對應")
    lines.append("")
    lines.append("> 依前例功能名稱比對 `code_e2e_chain`（ORW訂單系統）及 `code_menu`（MBMS）推算候選頁面與程式鏈。")
    lines.append("> **請人工核對是否為本次 SCN 真正異動的頁面，並補充「修改後業務邏輯」欄。**")
    lines.append("")

    if page_map:
        lines.append(
            "| 系統 | 功能名稱（前例）| 頁面名稱 | 作業代碼 | BackingBean | BD 層 | 配合修改方法 |"
        )
        lines.append("|---|---|---|---|---|---|---|")
        rendered_codes: set = set()
        for fname, pages in page_map.items():
            for pg in pages:
                bc = pg.get("base_code", "")
                src = pg.get("source", "?")
                # ORW entries carry chain data directly; MBMS entries don't
                bb = pg.get("backing_bean", "") or "—"
                delegates = pg.get("delegates") or []
                bds = ", ".join(delegates[:2]) or "—"
                # Build bd_methods from embedded bd_calls
                bd_calls = pg.get("bd_calls") or {}
                bd_methods: List[str] = []
                for bd in delegates[:2]:
                    for m in (bd_calls.get(bd) or [])[:3]:
                        bd_methods.append(f"{bd.split('.')[-1]}.{m}")
                methods = "; ".join(bd_methods[:3]) or "—"
                key = f"{src}|{fname[:15]}|{bc}"
                if key in rendered_codes:
                    continue
                rendered_codes.add(key)
                bb_short = bb.split(".")[-1] if bb != "—" else "—"
                lines.append(
                    f"| {src.upper()} | {_md_cell(fname[:20])} | {_md_cell(pg['menu_name'][:30])} "
                    f"| `{bc}` | {_md_cell(bb_short[:28])} "
                    f"| {_md_cell(bds[:35])} | {_md_cell(methods[:55])} |"
                )
        lines.append("")

        # Dedup detailed program chain (ORW entries only)
        detailed: Dict[str, Dict[str, Any]] = {}
        for pages in page_map.values():
            for pg in pages:
                if pg.get("source") == "orw":
                    bc = pg.get("base_code", "")
                    if bc and bc not in detailed:
                        delegates = pg.get("delegates") or []
                        bd_calls = pg.get("bd_calls") or {}
                        bd_methods_full: List[str] = []
                        for bd in delegates[:3]:
                            for m in (bd_calls.get(bd) or [])[:4]:
                                bd_methods_full.append(f"{bd}.{m}")
                        detailed[bc] = {
                            "menu_name": pg.get("menu_name", ""),
                            "backing_bean": pg.get("backing_bean", ""),
                            "delegates": delegates,
                            "bd_methods": bd_methods_full[:6],
                            "action_methods": pg.get("action_methods") or [],
                        }

        if detailed:
            lines.append("### 配合修改程式明細（ORW 訂單系統，依作業代碼）")
            lines.append("")
            for bc, ch in detailed.items():
                lines.append(f"#### `{bc}`（{ch['menu_name']}）")
                bb = ch.get("backing_bean", "")
                if bb:
                    lines.append(f"- **BackingBean**：`{bb.split('.')[-1]}`")
                dlist = ch.get("delegates", [])
                if dlist:
                    lines.append(f"- **BD 層**：{', '.join(f'`{d}`' for d in dlist)}")
                for m in ch.get("bd_methods", []):
                    lines.append(f"  - `{m}`")
                actions = ch.get("action_methods", [])
                if actions:
                    lines.append(f"- **Action 方法**：{', '.join(f'`{m}`' for m in actions)}")
                lines.append("- **修改後業務邏輯**：（待 BA 填入）")
                lines.append("")

    else:
        lines.append("_（前例功能名稱無法比對到 code_menu 頁面；請以作業代碼手工查詢）_")
        lines.append("")

    # 候選測試案例 — from the closest precedents, to seed this SCN's verification.
    tc_detail_n = min(3, len(precedents))
    tc_body: List[str] = []
    tc_total = 0
    for p in precedents[:tc_detail_n]:
        tcs = p.get("test_cases", []) or []
        if not tcs:
            continue
        tc_body.append(f"#### {_strip_swds_prefix(p['title'])}（相似度 {p['score']}）")
        tc_body.append("")
        tc_body.append("| 案例名稱 | 前提條件 | 預期結果 |")
        tc_body.append("|---|---|---|")
        for c in tcs[:8]:
            tc_body.append(
                f"| {_md_cell(c.get('name', ''))} | {_md_cell(c.get('precondition', ''))} "
                f"| {_md_cell(c.get('expected', ''))} |"
            )
            tc_total += 1
        tc_body.append("")
    lines.append(f"## 候選測試案例（取前 {tc_detail_n} 份前例）")
    lines.append("")
    if tc_body:
        lines.append("> 下列為相似前例的測試案例，供設計本次驗證情境參考；請依實際 SCN 增刪與調整門號前提。")
        lines.append("")
        lines.extend(tc_body)
    else:
        lines.append("_（前例未偵測到測試案例）_")
        lines.append("")

    # 撰寫提示
    lines.append("## 撰寫提示")
    lines.append("")
    lines.append("1. 上表/清單為「候選」而非結論：逐項對照本次 SCN，刪去不相關者、補上前例未涵蓋者。")
    lines.append("2. 第2章 前例數 ≥ 2、或第3章「跨前例重複出現」的項目為高機率要動者，優先確認。")
    lines.append("3. 欄位定義請開對應來源 SWDS 逐欄比對（本草稿只帶參考欄位數，不帶欄位明細）。")
    lines.append("4. 候選測試案例需覆蓋本次 SCN 的正常/例外情境，並補齊門號前提與預期結果。")
    lines.append("5. 選單路徑箭頭（→）在部分來源為符號字型、抽取時可能遺失，請以來源文件為準。")
    lines.append("6. 「頁面名稱 ↔ 程式鏈」為 KB 比對結果（非精確）：請確認 BackingBean/BD 是否為本次真正異動目標，並填入修改後業務邏輯。")
    lines.append("")

    # 追溯資訊（待填）
    lines.append("## 追溯資訊（待填）")
    lines.append("")
    lines.append("- 需求編號：（待填）")
    lines.append("- MR編號：（待填）")
    lines.append("")

    # 參考前例
    lines.append("## 參考前例")
    lines.append("")
    for p in precedents:
        lines.append(f"- [{p['score']}] {p['title']}（{p['category']}）")
    lines.append("")

    markdown = "\n".join(lines)
    return {
        "query": text,
        "precedent_count": len(precedents),
        "candidate_table_count": len(table_rows),
        "candidate_function_count": len(func_rows),
        "candidate_test_case_count": tc_total,
        "recurring_function_count": len(recurring),
        "resolved_page_count": sum(len(v) for v in page_map.values()),
        "resolved_program_count": len(chain_map),
        "high_confidence_tables": [name for freq, name, _ in table_rows if freq >= 2],
        "draft_markdown": markdown,
        "precedents": [
            {k: v for k, v in p.items() if k not in ("functions", "test_cases")}
            for p in precedents
        ],
    }
